#!/usr/bin/env python3
"""
Enhanced Word to PDF Page Structure Preservation Testing
Tests multi-page document conversion, page break preservation, and page count verification
"""

import requests
import io
import os
import tempfile
import zipfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from pypdf import PdfReader
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Get backend URL from environment
BACKEND_URL = "https://doctools-6.preview.emergentagent.com/api"

def create_multi_page_docx(pages: int = 4) -> tuple:
    """Create a test DOCX file with content spanning multiple pages"""
    doc = Document()
    
    # Add title page content
    doc.add_heading('Multi-Page Test Document', 0)
    doc.add_paragraph('This document is designed to test page structure preservation in PDF conversion.')
    doc.add_paragraph('Document created for testing purposes with controlled page content.')
    
    # Add substantial content for each page
    for page_num in range(1, pages + 1):
        # Add page break before each new page (except first)
        if page_num > 1:
            doc.add_page_break()
        
        # Add page header
        doc.add_heading(f'Page {page_num} Content', level=1)
        
        # Add multiple paragraphs to fill the page
        for para_num in range(1, 8):  # 7 paragraphs per page
            content = f"This is paragraph {para_num} on page {page_num}. " * 10
            content += f"This content is specifically designed to test page structure preservation. "
            content += f"Each paragraph contains substantial text to ensure proper page flow and height estimation. "
            content += f"Page {page_num}, Paragraph {para_num} marker for verification."
            doc.add_paragraph(content)
        
        # Add a table on every other page
        if page_num % 2 == 0:
            table = doc.add_table(rows=4, cols=3)
            table.style = 'Table Grid'
            
            # Add table headers
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = f'Page {page_num} Col 1'
            hdr_cells[1].text = f'Page {page_num} Col 2'
            hdr_cells[2].text = f'Page {page_num} Col 3'
            
            # Add table data
            for row_idx in range(1, 4):
                row_cells = table.rows[row_idx].cells
                row_cells[0].text = f'Row {row_idx} Data A'
                row_cells[1].text = f'Row {row_idx} Data B'
                row_cells[2].text = f'Row {row_idx} Data C'
        
        # Add spacing content
        doc.add_paragraph('')  # Empty paragraph for spacing
        doc.add_paragraph(f'End of page {page_num} content. This paragraph marks the conclusion of page {page_num}.')
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue(), pages

def create_explicit_page_break_docx() -> bytes:
    """Create a DOCX with explicit page breaks"""
    doc = Document()
    
    # Page 1
    doc.add_heading('Document with Explicit Page Breaks', 0)
    doc.add_paragraph('This is the first page content.')
    for i in range(5):
        doc.add_paragraph(f'First page paragraph {i+1} with substantial content to fill the page properly.')
    
    # Explicit page break
    doc.add_page_break()
    
    # Page 2
    doc.add_heading('Second Page After Explicit Break', level=1)
    doc.add_paragraph('This content should appear on the second page after an explicit page break.')
    for i in range(6):
        doc.add_paragraph(f'Second page paragraph {i+1} with content to test page break preservation.')
    
    # Another explicit page break
    doc.add_page_break()
    
    # Page 3
    doc.add_heading('Third Page After Another Break', level=1)
    doc.add_paragraph('This content should appear on the third page.')
    for i in range(4):
        doc.add_paragraph(f'Third page paragraph {i+1} to verify explicit page breaks work correctly.')
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()

def get_pdf_page_count(pdf_content: bytes) -> int:
    """Extract page count from PDF content"""
    try:
        pdf_buffer = io.BytesIO(pdf_content)
        reader = PdfReader(pdf_buffer)
        return len(reader.pages)
    except Exception as e:
        logger.error(f"Error reading PDF page count: {e}")
        return 0

def extract_pdf_text(pdf_content: bytes) -> str:
    """Extract text content from PDF for verification"""
    try:
        pdf_buffer = io.BytesIO(pdf_content)
        reader = PdfReader(pdf_buffer)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF text: {e}")
        return ""

def test_multi_page_conversion():
    """Test multi-page document conversion with page count verification"""
    print("\n🔍 Testing Multi-Page Document Conversion...")
    
    test_cases = [
        {"pages": 2, "name": "2-page document"},
        {"pages": 4, "name": "4-page document"},
        {"pages": 6, "name": "6-page document"}
    ]
    
    results = []
    
    for test_case in test_cases:
        print(f"\n   Testing {test_case['name']}...")
        
        try:
            # Create test document
            docx_content, expected_pages = create_multi_page_docx(test_case['pages'])
            
            # Test single_pdf=true
            files = {'files': (f'test_{test_case["pages"]}page.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
            data = {'single_pdf': 'true'}
            
            response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
            
            if response.status_code == 200:
                pdf_page_count = get_pdf_page_count(response.content)
                pdf_text = extract_pdf_text(response.content)
                
                print(f"   Expected pages: {expected_pages}")
                print(f"   PDF pages: {pdf_page_count}")
                print(f"   PDF size: {len(response.content)} bytes")
                
                # Verify page count is reasonable (allow some variance due to content flow)
                page_count_ok = abs(pdf_page_count - expected_pages) <= 1
                
                # Verify content preservation
                content_markers = [f"Page {i} Content" for i in range(1, expected_pages + 1)]
                content_preserved = all(marker in pdf_text for marker in content_markers)
                
                # Verify paragraph markers
                paragraph_markers = [f"Page {i}, Paragraph" for i in range(1, expected_pages + 1)]
                paragraphs_preserved = all(any(marker in pdf_text for marker in paragraph_markers) for i in range(1, expected_pages + 1))
                
                if page_count_ok and content_preserved and paragraphs_preserved:
                    print(f"   ✅ {test_case['name']} conversion successful")
                    print(f"      - Page count preserved (±1): {pdf_page_count}/{expected_pages}")
                    print(f"      - Content markers found: {len([m for m in content_markers if m in pdf_text])}/{len(content_markers)}")
                    results.append(True)
                else:
                    print(f"   ❌ {test_case['name']} conversion issues:")
                    if not page_count_ok:
                        print(f"      - Page count mismatch: {pdf_page_count} vs expected {expected_pages}")
                    if not content_preserved:
                        print(f"      - Content markers missing: {[m for m in content_markers if m not in pdf_text]}")
                    if not paragraphs_preserved:
                        print(f"      - Paragraph markers missing")
                    results.append(False)
            else:
                print(f"   ❌ {test_case['name']} conversion failed: {response.status_code}")
                results.append(False)
                
        except Exception as e:
            print(f"   ❌ {test_case['name']} test error: {str(e)}")
            results.append(False)
    
    return all(results)

def test_explicit_page_breaks():
    """Test preservation of explicit page breaks"""
    print("\n🔍 Testing Explicit Page Break Preservation...")
    
    try:
        # Create document with explicit page breaks
        docx_content = create_explicit_page_break_docx()
        
        files = {'files': ('explicit_breaks.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code == 200:
            pdf_page_count = get_pdf_page_count(response.content)
            pdf_text = extract_pdf_text(response.content)
            
            print(f"   PDF pages: {pdf_page_count}")
            print(f"   PDF size: {len(response.content)} bytes")
            
            # Should have at least 3 pages due to explicit breaks
            expected_min_pages = 3
            page_count_ok = pdf_page_count >= expected_min_pages
            
            # Verify content from each page is present
            page_markers = [
                "Document with Explicit Page Breaks",
                "Second Page After Explicit Break", 
                "Third Page After Another Break"
            ]
            content_preserved = all(marker in pdf_text for marker in page_markers)
            
            if page_count_ok and content_preserved:
                print(f"   ✅ Explicit page breaks preserved successfully")
                print(f"      - Page count: {pdf_page_count} (expected ≥{expected_min_pages})")
                print(f"      - All page markers found: {len([m for m in page_markers if m in pdf_text])}/{len(page_markers)}")
                return True
            else:
                print(f"   ❌ Explicit page break preservation issues:")
                if not page_count_ok:
                    print(f"      - Insufficient pages: {pdf_page_count} (expected ≥{expected_min_pages})")
                if not content_preserved:
                    print(f"      - Missing page markers: {[m for m in page_markers if m not in pdf_text]}")
                return False
        else:
            print(f"   ❌ Explicit page break test failed: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Explicit page break test error: {str(e)}")
        return False

def test_content_distribution():
    """Test that content is properly distributed across pages"""
    print("\n🔍 Testing Content Distribution Across Pages...")
    
    try:
        # Create a document with known content distribution
        docx_content, expected_pages = create_multi_page_docx(4)
        
        files = {'files': ('distribution_test.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code == 200:
            pdf_page_count = get_pdf_page_count(response.content)
            pdf_text = extract_pdf_text(response.content)
            
            print(f"   PDF pages: {pdf_page_count}")
            print(f"   PDF size: {len(response.content)} bytes")
            print(f"   Text length: {len(pdf_text)} characters")
            
            # Verify content is not compressed into too few pages
            min_expected_pages = max(1, expected_pages - 1)  # Allow 1 page variance
            content_not_compressed = pdf_page_count >= min_expected_pages
            
            # Verify substantial content is present
            substantial_content = len(pdf_text) > 1000  # Should have substantial text
            
            # Verify table content is present (tables are on even pages)
            table_markers = ["Col 1", "Col 2", "Col 3", "Row 1 Data"]
            tables_preserved = any(marker in pdf_text for marker in table_markers)
            
            # Verify page-specific content distribution
            page_content_distributed = True
            for page_num in range(1, expected_pages + 1):
                page_marker = f"Page {page_num} Content"
                if page_marker not in pdf_text:
                    page_content_distributed = False
                    break
            
            if content_not_compressed and substantial_content and tables_preserved and page_content_distributed:
                print(f"   ✅ Content distribution preserved successfully")
                print(f"      - Pages not over-compressed: {pdf_page_count} ≥ {min_expected_pages}")
                print(f"      - Substantial content: {len(pdf_text)} characters")
                print(f"      - Tables preserved: {tables_preserved}")
                print(f"      - Page content distributed: {page_content_distributed}")
                return True
            else:
                print(f"   ❌ Content distribution issues:")
                if not content_not_compressed:
                    print(f"      - Content over-compressed: {pdf_page_count} < {min_expected_pages}")
                if not substantial_content:
                    print(f"      - Insufficient content: {len(pdf_text)} characters")
                if not tables_preserved:
                    print(f"      - Tables not preserved")
                if not page_content_distributed:
                    print(f"      - Page content not properly distributed")
                return False
        else:
            print(f"   ❌ Content distribution test failed: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Content distribution test error: {str(e)}")
        return False

def test_zip_mode_page_structure():
    """Test page structure preservation in ZIP mode"""
    print("\n🔍 Testing Page Structure in ZIP Mode...")
    
    try:
        # Create multiple documents with different page counts
        doc1_content, doc1_pages = create_multi_page_docx(2)
        doc2_content, doc2_pages = create_multi_page_docx(3)
        
        files = [
            ('files', ('doc1_2pages.docx', doc1_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')),
            ('files', ('doc2_3pages.docx', doc2_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        data = {'single_pdf': 'false'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code == 200:
            # Extract ZIP and analyze each PDF
            zip_buffer = io.BytesIO(response.content)
            with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
                pdf_files = zip_file.namelist()
                print(f"   ZIP contains: {pdf_files}")
                
                results = []
                for pdf_filename in pdf_files:
                    pdf_content = zip_file.read(pdf_filename)
                    pdf_page_count = get_pdf_page_count(pdf_content)
                    pdf_text = extract_pdf_text(pdf_content)
                    
                    print(f"   {pdf_filename}: {pdf_page_count} pages, {len(pdf_text)} chars")
                    
                    # Verify reasonable page count
                    expected_pages = 2 if "doc1" in pdf_filename else 3
                    page_count_ok = abs(pdf_page_count - expected_pages) <= 1
                    
                    # Verify content preservation
                    content_ok = len(pdf_text) > 500  # Substantial content
                    
                    results.append(page_count_ok and content_ok)
                
                if all(results) and len(pdf_files) == 2:
                    print(f"   ✅ ZIP mode page structure preserved successfully")
                    return True
                else:
                    print(f"   ❌ ZIP mode page structure issues")
                    return False
        else:
            print(f"   ❌ ZIP mode test failed: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ ZIP mode test error: {str(e)}")
        return False

def main():
    """Run all page structure preservation tests"""
    print("🚀 Starting Enhanced Word to PDF Page Structure Tests")
    print(f"Backend URL: {BACKEND_URL}")
    print("=" * 70)
    
    test_results = {}
    
    # Run all page structure tests
    test_results['multi_page_conversion'] = test_multi_page_conversion()
    test_results['explicit_page_breaks'] = test_explicit_page_breaks()
    test_results['content_distribution'] = test_content_distribution()
    test_results['zip_mode_page_structure'] = test_zip_mode_page_structure()
    
    # Summary
    print("\n" + "=" * 70)
    print("📊 PAGE STRUCTURE TEST SUMMARY")
    print("=" * 70)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{test_name.replace('_', ' ').title()}: {status}")
        if result:
            passed += 1
    
    print(f"\nOverall: {passed}/{total} page structure tests passed")
    
    if passed == total:
        print("🎉 All page structure tests passed! Enhanced Word to PDF conversion is working correctly.")
        return True
    else:
        print("⚠️  Some page structure tests failed. Please check the issues above.")
        return False

if __name__ == "__main__":
    main()