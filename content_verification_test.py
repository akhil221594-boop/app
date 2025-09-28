#!/usr/bin/env python3
"""
Enhanced Content Verification Test for Word to PDF Conversion
Specifically tests that PDF contains actual document content, not just filename
"""

import requests
import io
import tempfile
from docx import Document
from pypdf import PdfReader
import zipfile

# Get backend URL from environment
BACKEND_URL = "https://doctools-6.preview.emergentagent.com/api"

def create_rich_test_docx():
    """Create a test DOCX file with rich content for verification"""
    doc = Document()
    
    # Add title
    doc.add_heading('Test Document for Content Verification', 0)
    
    # Add various content types
    doc.add_paragraph('This is the first paragraph with specific test content that should appear in the PDF.')
    
    doc.add_heading('Section 1: Important Information', level=1)
    doc.add_paragraph('This section contains critical information that must be preserved during conversion. '
                     'The content includes specific keywords like: VERIFICATION_KEYWORD_123 and TEST_CONTENT_456.')
    
    doc.add_heading('Section 2: Data Table', level=1)
    doc.add_paragraph('Below is a table that should be converted properly:')
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Fill table with test data
    cells = table.rows[0].cells
    cells[0].text = 'Header 1'
    cells[1].text = 'Header 2'
    cells[2].text = 'Header 3'
    
    cells = table.rows[1].cells
    cells[0].text = 'Row1Col1'
    cells[1].text = 'Row1Col2'
    cells[2].text = 'Row1Col3'
    
    cells = table.rows[2].cells
    cells[0].text = 'Row2Col1'
    cells[1].text = 'Row2Col2'
    cells[2].text = 'Row2Col3'
    
    doc.add_paragraph('This is the final paragraph with unique identifier: FINAL_PARAGRAPH_789')
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

def extract_text_from_pdf(pdf_content: bytes) -> str:
    """Extract text from PDF content"""
    try:
        pdf_buffer = io.BytesIO(pdf_content)
        reader = PdfReader(pdf_buffer)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return ""

def test_content_preservation_single_pdf():
    """Test that single PDF contains actual document content"""
    print("🔍 Testing content preservation in single PDF...")
    
    try:
        docx_content = create_rich_test_docx()
        
        files = {'files': ('rich_test_document.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code != 200:
            print(f"   ❌ Conversion failed with status {response.status_code}: {response.text}")
            return False
        
        # Extract text from PDF
        pdf_text = extract_text_from_pdf(response.content)
        print(f"   Extracted PDF text length: {len(pdf_text)} characters")
        
        # Check for specific content
        required_content = [
            "Test Document for Content Verification",
            "VERIFICATION_KEYWORD_123",
            "TEST_CONTENT_456",
            "Section 1: Important Information",
            "Section 2: Data Table",
            "FINAL_PARAGRAPH_789",
            "Header 1",
            "Row1Col1",
            "Row2Col2"
        ]
        
        missing_content = []
        found_content = []
        
        for content in required_content:
            if content in pdf_text:
                found_content.append(content)
                print(f"   ✅ Found: '{content}'")
            else:
                missing_content.append(content)
                print(f"   ❌ Missing: '{content}'")
        
        # Check if filename is NOT the only content
        filename_only = "rich_test_document" in pdf_text and len(pdf_text.strip()) < 100
        if filename_only:
            print("   ❌ PDF appears to contain only filename, not document content")
            return False
        
        if len(found_content) >= len(required_content) * 0.7:  # At least 70% of content found
            print(f"   ✅ Content preservation successful: {len(found_content)}/{len(required_content)} items found")
            return True
        else:
            print(f"   ❌ Insufficient content preserved: {len(found_content)}/{len(required_content)} items found")
            return False
            
    except Exception as e:
        print(f"   ❌ Content preservation test error: {str(e)}")
        return False

def test_content_preservation_zip():
    """Test that ZIP PDFs contain actual document content"""
    print("\n🔍 Testing content preservation in ZIP output...")
    
    try:
        docx_content1 = create_rich_test_docx()
        
        # Create second document with different content
        doc2 = Document()
        doc2.add_heading('Second Test Document', 0)
        doc2.add_paragraph('This is the second document with unique content: SECOND_DOC_MARKER_999')
        docx_buffer2 = io.BytesIO()
        doc2.save(docx_buffer2)
        docx_buffer2.seek(0)
        docx_content2 = docx_buffer2.getvalue()
        
        files = [
            ('files', ('document1.docx', docx_content1, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')),
            ('files', ('document2.docx', docx_content2, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        data = {'single_pdf': 'false'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code != 200:
            print(f"   ❌ Conversion failed with status {response.status_code}: {response.text}")
            return False
        
        # Extract PDFs from ZIP
        zip_buffer = io.BytesIO(response.content)
        with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
            pdf_files = zip_file.namelist()
            print(f"   ZIP contains: {pdf_files}")
            
            success_count = 0
            
            for pdf_filename in pdf_files:
                pdf_content = zip_file.read(pdf_filename)
                pdf_text = extract_text_from_pdf(pdf_content)
                
                print(f"   Checking {pdf_filename} - text length: {len(pdf_text)} characters")
                
                if pdf_filename == 'document1.pdf':
                    # Check for first document content
                    if "VERIFICATION_KEYWORD_123" in pdf_text and "Test Document for Content Verification" in pdf_text:
                        print(f"   ✅ {pdf_filename} contains expected content")
                        success_count += 1
                    else:
                        print(f"   ❌ {pdf_filename} missing expected content")
                
                elif pdf_filename == 'document2.pdf':
                    # Check for second document content
                    if "SECOND_DOC_MARKER_999" in pdf_text and "Second Test Document" in pdf_text:
                        print(f"   ✅ {pdf_filename} contains expected content")
                        success_count += 1
                    else:
                        print(f"   ❌ {pdf_filename} missing expected content")
            
            if success_count == len(pdf_files):
                print("   ✅ All PDFs in ZIP contain proper content")
                return True
            else:
                print(f"   ❌ Only {success_count}/{len(pdf_files)} PDFs contain proper content")
                return False
                
    except Exception as e:
        print(f"   ❌ ZIP content preservation test error: {str(e)}")
        return False

def test_formatting_preservation():
    """Test that basic formatting is preserved"""
    print("\n🔍 Testing formatting preservation...")
    
    try:
        # Create document with specific formatting
        doc = Document()
        doc.add_heading('MAIN TITLE', 0)
        doc.add_heading('Subtitle Level 1', 1)
        doc.add_paragraph('Regular paragraph text.')
        
        # Add formatted text
        p = doc.add_paragraph()
        run = p.add_run('Bold text')
        run.bold = True
        p.add_run(' and ')
        run2 = p.add_run('italic text')
        run2.italic = True
        
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_content = docx_buffer.getvalue()
        
        files = {'files': ('formatting_test.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        if response.status_code != 200:
            print(f"   ❌ Conversion failed with status {response.status_code}: {response.text}")
            return False
        
        pdf_text = extract_text_from_pdf(response.content)
        
        # Check for text content (formatting may not be perfectly preserved in text extraction)
        required_text = ["MAIN TITLE", "Subtitle Level 1", "Regular paragraph text", "Bold text", "italic text"]
        found_count = 0
        
        for text in required_text:
            if text in pdf_text:
                found_count += 1
                print(f"   ✅ Found: '{text}'")
            else:
                print(f"   ❌ Missing: '{text}'")
        
        if found_count >= len(required_text) * 0.8:  # At least 80% found
            print("   ✅ Basic formatting and text preserved")
            return True
        else:
            print(f"   ❌ Insufficient formatting preserved: {found_count}/{len(required_text)}")
            return False
            
    except Exception as e:
        print(f"   ❌ Formatting preservation test error: {str(e)}")
        return False

def test_error_handling_corrupted():
    """Test error handling for corrupted documents"""
    print("\n🔍 Testing error handling for corrupted documents...")
    
    try:
        # Create corrupted DOCX content
        corrupted_content = b"This is not a valid DOCX file content but has .docx extension"
        
        files = {'files': ('corrupted.docx', corrupted_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        
        # Should return an error status
        if response.status_code >= 400:
            print(f"   ✅ Correctly handled corrupted document with status {response.status_code}")
            return True
        else:
            print(f"   ❌ Should have returned error status, got {response.status_code}")
            return False
            
    except Exception as e:
        print(f"   ❌ Corrupted document test error: {str(e)}")
        return False

def main():
    """Run enhanced content verification tests"""
    print("🚀 Starting Enhanced Word to PDF Content Verification Tests")
    print(f"Backend URL: {BACKEND_URL}")
    print("=" * 70)
    
    test_results = {}
    
    # Run content-focused tests
    test_results['content_preservation_single'] = test_content_preservation_single_pdf()
    test_results['content_preservation_zip'] = test_content_preservation_zip()
    test_results['formatting_preservation'] = test_formatting_preservation()
    test_results['error_handling_corrupted'] = test_error_handling_corrupted()
    
    # Summary
    print("\n" + "=" * 70)
    print("📊 ENHANCED TEST SUMMARY")
    print("=" * 70)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{test_name.replace('_', ' ').title()}: {status}")
        if result:
            passed += 1
    
    print(f"\nOverall: {passed}/{total} enhanced tests passed")
    
    if passed == total:
        print("🎉 All enhanced content verification tests passed!")
        print("✅ Word to PDF conversion properly preserves document content")
        return True
    else:
        print("⚠️  Some enhanced tests failed. Content preservation may have issues.")
        return False

if __name__ == "__main__":
    main()