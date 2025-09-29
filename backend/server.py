from fastapi import FastAPI, APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field
from typing import List
import uuid
from datetime import datetime
import io
import zipfile
import tempfile
from docx import Document
import docx2txt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.utils import ImageReader
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
import PyPDF2
from PIL import Image as PILImage
import xml.etree.ElementTree as ET


ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")


# Define Models
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

# Add your routes to the router instead of directly to app
@api_router.get("/")
async def root():
    return {"message": "Hello World"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.dict()
    status_obj = StatusCheck(**status_dict)
    _ = await db.status_checks.insert_one(status_obj.dict())
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    status_checks = await db.status_checks.find().to_list(1000)
    return [StatusCheck(**status_check) for status_check in status_checks]

# PDF Processing Functions
def convert_docx_to_pdf(docx_file_content: bytes, filename: str) -> bytes:
    """Convert a Word document to PDF preserving original pagination"""
    try:
        # Create Document from bytes
        doc = Document(io.BytesIO(docx_file_content))
        
        # Create PDF buffer
        pdf_buffer = io.BytesIO()
        
        # Create PDF document with Word-like page settings
        pdf_doc = SimpleDocTemplate(
            pdf_buffer, 
            pagesize=A4,  # Use A4 which is more common for Word docs
            rightMargin=1*inch,
            leftMargin=1*inch,
            topMargin=1*inch,
            bottomMargin=1*inch
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Process each section which often corresponds to page breaks in Word
        sections = doc.sections
        for section_idx, section in enumerate(sections):
            if section_idx > 0:  # Add page break between sections
                story.append(PageBreak())
            
            # Process paragraphs within each section
            section_paragraphs = []
            
            # Collect paragraphs for this section
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    section_paragraphs.append(paragraph)
        
        # Enhanced paragraph processing with page break detection
        page_height_used = 0
        max_page_height = 9 * inch  # Approximate text area height
        
        for i, paragraph in enumerate(doc.paragraphs):
            if not paragraph.text.strip():
                continue
                
            # Check for explicit page breaks in Word XML
            paragraph_xml = paragraph._element.xml
            has_page_break = False
            
            # Look for page break elements in Word XML
            if 'w:br' in paragraph_xml and 'w:type="page"' in paragraph_xml:
                has_page_break = True
            
            # Check if this paragraph starts on a new page based on Word's page break settings
            if has_page_break or (hasattr(paragraph, 'page_break_before') and paragraph.page_break_before):
                story.append(PageBreak())
                page_height_used = 0
            
            # Determine paragraph style and formatting
            text_content = paragraph.text
            
            # Handle different paragraph styles
            if paragraph.style.name.startswith('Heading'):
                level = 1
                try:
                    if paragraph.style.name[-1].isdigit():
                        level = int(paragraph.style.name[-1])
                except:
                    pass
                
                heading_style = ParagraphStyle(
                    f'Heading{level}',
                    parent=styles[f'Heading{min(level, 3)}'],
                    fontSize=max(14, 18 - level * 2),
                    spaceAfter=12,
                    spaceBefore=12,
                    keepWithNext=True  # Keep headings with following content
                )
                p = Paragraph(text_content, heading_style)
                page_height_used += 40  # Approximate height for headings
                
            else:
                # Normal paragraph
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=11,
                    spaceAfter=6,
                    spaceBefore=6,
                    alignment=TA_LEFT
                )
                p = Paragraph(text_content, normal_style)
                # Estimate paragraph height (rough calculation)
                lines = len(text_content) // 80 + 1  # Approximate characters per line
                page_height_used += lines * 14  # Approximate line height
            
            story.append(p)
            
            # Insert page break if we're approaching page limit and this isn't the last paragraph
            if page_height_used > max_page_height * 0.85 and i < len(doc.paragraphs) - 1:
                next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
                if next_para and next_para.text.strip():
                    story.append(PageBreak())
                    page_height_used = 0
        
        # Handle images with proper sizing and placement
        try:
            image_elements = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image_stream = io.BytesIO(image_data)
                        
                        # Create properly sized image
                        img = Image(image_stream)
                        
                        # Scale image to fit page properly
                        max_width = 6 * inch
                        max_height = 7 * inch
                        
                        if hasattr(img, 'imageWidth') and hasattr(img, 'imageHeight'):
                            scale_w = max_width / img.imageWidth
                            scale_h = max_height / img.imageHeight
                            scale = min(scale_w, scale_h, 1)
                            
                            img.drawWidth = img.imageWidth * scale
                            img.drawHeight = img.imageHeight * scale
                        else:
                            img.drawWidth = 4 * inch
                            img.drawHeight = 3 * inch
                        
                        image_elements.append(img)
                    except Exception as e:
                        logging.warning(f"Could not process image: {e}")
            
            # Add images at the end or integrate them properly if needed
            for img in image_elements[:3]:  # Limit to 3 images to prevent issues
                story.append(Spacer(1, 12))
                story.append(img)
                story.append(Spacer(1, 12))
                
        except Exception as e:
            logging.warning(f"Error processing images: {e}")
        
        # Build PDF with proper page management
        pdf_doc.build(story)
        pdf_buffer.seek(0)
        
        return pdf_buffer.getvalue()
        
    except Exception as e:
        logging.error(f"Error converting DOCX to PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Error converting document: {str(e)}")

def compress_pdf(pdf_content: bytes) -> bytes:
    """Compress PDF to reduce file size"""
    try:
        input_pdf = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        output_pdf = PyPDF2.PdfWriter()
        
        # Copy all pages while applying compression
        for page in input_pdf.pages:
            # Compress page content
            page.compress_content_streams()
            output_pdf.add_page(page)
        
        # Create compressed PDF buffer
        compressed_buffer = io.BytesIO()
        output_pdf.write(compressed_buffer)
        compressed_buffer.seek(0)
        
        return compressed_buffer.getvalue()
        
    except Exception as e:
        logging.error(f"Error compressing PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Error compressing PDF: {str(e)}")

# New API Endpoints for PDF Processing
@api_router.post("/convert-word-to-pdf")
async def convert_word_to_pdf(
    files: List[UploadFile] = File(...),
    output_type: str = Form("single")
):
    """Convert Word documents to PDF"""
    try:
        if not files:
            raise HTTPException(status_code=400, detail="No files provided")
        
        # Validate file types
        for file in files:
            if not file.filename.lower().endswith(('.docx', '.doc')):
                raise HTTPException(status_code=400, detail=f"Invalid file type: {file.filename}")
        
        converted_pdfs = []
        
        # Process each file
        for file in files:
            content = await file.read()
            
            # Convert to PDF
            if file.filename.lower().endswith('.docx'):
                pdf_content = convert_docx_to_pdf(content, file.filename)
                pdf_filename = file.filename.rsplit('.', 1)[0] + '.pdf'
                converted_pdfs.append((pdf_filename, pdf_content))
            else:
                raise HTTPException(status_code=400, detail="Only .docx files are currently supported")
        
        if output_type == "single" and len(converted_pdfs) > 1:
            # Merge multiple PDFs into one
            merged_buffer = io.BytesIO()
            pdf_writer = PyPDF2.PdfWriter()
            
            for _, pdf_content in converted_pdfs:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
            
            pdf_writer.write(merged_buffer)
            merged_buffer.seek(0)
            
            return StreamingResponse(
                io.BytesIO(merged_buffer.getvalue()),
                media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=converted-documents.pdf"}
            )
        
        elif output_type == "multiple" or len(converted_pdfs) > 1:
            # Create ZIP file with multiple PDFs
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for pdf_filename, pdf_content in converted_pdfs:
                    zip_file.writestr(pdf_filename, pdf_content)
            
            zip_buffer.seek(0)
            
            return StreamingResponse(
                io.BytesIO(zip_buffer.getvalue()),
                media_type="application/zip",
                headers={"Content-Disposition": "attachment; filename=converted-documents.zip"}
            )
        
        else:
            # Single file
            pdf_filename, pdf_content = converted_pdfs[0]
            
            return StreamingResponse(
                io.BytesIO(pdf_content),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={pdf_filename}"}
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in convert_word_to_pdf: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@api_router.post("/compress-pdf")
async def compress_pdf_endpoint(file: UploadFile = File(...)):
    """Compress PDF file to reduce size"""
    try:
        if not file.filename.lower().endswith('.pdf'):
            raise HTTPException(status_code=400, detail="File must be a PDF")
        
        # Read PDF content
        pdf_content = await file.read()
        
        # Compress the PDF
        compressed_content = compress_pdf(pdf_content)
        
        # Create filename
        base_name = file.filename.rsplit('.', 1)[0]
        compressed_filename = f"compressed_{base_name}.pdf"
        
        return StreamingResponse(
            io.BytesIO(compressed_content),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={compressed_filename}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Error in compress_pdf_endpoint: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()
