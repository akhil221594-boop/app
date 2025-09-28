import io
import os
import zipfile
import tempfile
from typing import List
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.units import inch
from pypdf import PdfWriter, PdfReader
from PIL import Image
import logging
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class PDFConverter:
    
    @staticmethod
    def detect_page_breaks(docx_path: str) -> List[int]:
        """Detect explicit page breaks in the Word document"""
        page_break_indices = []
        try:
            import zipfile as zf
            with zf.ZipFile(docx_path, 'r') as docx_zip:
                document_xml = docx_zip.read('word/document.xml').decode('utf-8')
                
                # Count paragraphs and find page breaks
                paragraph_index = 0
                lines = document_xml.split('<w:p')
                
                for i, line in enumerate(lines):
                    if 'w:br' in line and ('w:type="page"' in line or 'type="page"' in line):
                        page_break_indices.append(paragraph_index)
                    if '<w:p' in line or i == 0:  # Count paragraphs
                        paragraph_index += 1
                        
        except Exception as e:
            logger.warning(f"Could not detect page breaks: {e}")
            
        return page_break_indices
    
    @staticmethod
    def estimate_content_height(text: str, style) -> float:
        """Estimate the height a text block will take in the PDF"""
        # Rough estimation: each line is about 14 points, each page is about 720 points usable
        lines = len(text.split('\n')) + (len(text) // 80)  # Assume 80 chars per line
        height = lines * (style.leading if hasattr(style, 'leading') else 14)
        return height
    
    @staticmethod
    def extract_images_from_docx(docx_path: str) -> List[tuple]:
        """Extract images from DOCX file"""
        images = []
        try:
            import zipfile as zf
            with zf.ZipFile(docx_path, 'r') as docx_zip:
                # Get all image files from the docx
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        img_data = docx_zip.read(img_file)
                        # Save to temp file
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_img:
                            temp_img.write(img_data)
                            images.append((temp_img.name, img_data))
                    except Exception as e:
                        logger.warning(f"Could not extract image {img_file}: {e}")
                        
        except Exception as e:
            logger.warning(f"Could not extract images: {e}")
            
        return images
    
    @staticmethod
    def docx_to_pdf(docx_file_content: bytes, filename: str) -> bytes:
        """Convert a DOCX file to PDF preserving original page structure and count"""
        temp_docx_path = None
        temp_images = []
        
        try:
            # Create a temporary file for the DOCX content
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx:
                temp_docx.write(docx_file_content)
                temp_docx_path = temp_docx.name

            # Create PDF buffer
            pdf_buffer = io.BytesIO()
            
            # Create PDF document with proper page settings to maintain original structure
            doc = SimpleDocTemplate(
                pdf_buffer, 
                pagesize=A4,
                rightMargin=72, leftMargin=72,  # Standard 1 inch margins
                topMargin=72, bottomMargin=72,
                title=filename.replace('.docx', ''),
                showBoundary=0  # Don't show page boundaries
            )
            
            # Parse DOCX document
            docx_doc = Document(temp_docx_path)
            story = []
            
            # Get styles with proper spacing for page flow
            styles = getSampleStyleSheet()
            
            # Create styles that maintain original document spacing
            normal_style = ParagraphStyle(
                'EnhancedNormal',
                parent=styles['Normal'],
                fontSize=11,
                leading=14,  # Line height
                spaceAfter=6,  # Space after paragraph (reduced to maintain flow)
                spaceBefore=3,  # Space before paragraph
                leftIndent=0,
                rightIndent=0,
                firstLineIndent=0,
                alignment=0,  # Left alignment
                bulletIndent=0,
                wordWrap='LTR'
            )
            
            heading1_style = ParagraphStyle(
                'EnhancedHeading1',
                parent=styles['Heading1'],
                fontSize=16,
                leading=20,
                spaceAfter=12,
                spaceBefore=12,
                textColor=colors.black,
                fontName='Helvetica-Bold'
            )
            
            heading2_style = ParagraphStyle(
                'EnhancedHeading2',
                parent=styles['Heading2'],
                fontSize=14,
                leading=18,
                spaceAfter=10,
                spaceBefore=10,
                textColor=colors.black,
                fontName='Helvetica-Bold'
            )
            
            # Extract images and detect page breaks
            extracted_images = PDFConverter.extract_images_from_docx(temp_docx_path)
            temp_images = [img_path for img_path, _ in extracted_images]
            page_breaks = PDFConverter.detect_page_breaks(temp_docx_path)
            
            image_index = 0
            paragraph_count = 0
            current_page_height = 0
            max_page_height = 650  # Approximate usable page height in points
            
            # Process paragraphs in order to maintain document flow
            for paragraph_idx, paragraph in enumerate(docx_doc.paragraphs):
                text = paragraph.text.strip()
                
                # Check if this paragraph should trigger a page break
                # Look for explicit page breaks in the paragraph
                if hasattr(paragraph, '_element') and paragraph._element.xml:
                    if 'w:br' in paragraph._element.xml and 'type="page"' in paragraph._element.xml:
                        from reportlab.platypus import PageBreak
                        story.append(PageBreak())
                
                if text:
                    paragraph_count += 1
                    
                    # Determine paragraph style based on Word formatting
                    selected_style = normal_style
                    
                    if paragraph.style.name.startswith('Heading'):
                        if '1' in paragraph.style.name:
                            selected_style = heading1_style
                        else:
                            selected_style = heading2_style
                    elif any(run.bold for run in paragraph.runs if run.text.strip()):
                        # If the paragraph has bold text, treat as heading
                        selected_style = heading2_style
                    
                    # Process runs for formatting
                    formatted_text = ""
                    for run in paragraph.runs:
                        run_text = run.text
                        if run_text:
                            if run.bold:
                                run_text = f"<b>{run_text}</b>"
                            if run.italic:
                                run_text = f"<i>{run_text}</i>"
                            formatted_text += run_text
                    
                    # If no formatting found, use plain text
                    if not formatted_text.strip():
                        formatted_text = text
                    
                    # Replace line breaks with proper PDF line breaks
                    formatted_text = formatted_text.replace('\n', '<br/>')
                    
                    para = Paragraph(formatted_text, selected_style)
                    story.append(para)
                    
                    # Add appropriate spacing based on content
                    if paragraph.style.name.startswith('Heading'):
                        story.append(Spacer(1, 8))
                    else:
                        story.append(Spacer(1, 3))
                    
                    # Insert images near related text (every few paragraphs)
                    if image_index < len(extracted_images) and paragraph_count % 3 == 0:
                        try:
                            img_path, _ = extracted_images[image_index]
                            # Calculate image size to fit within page margins
                            img = RLImage(img_path, width=5*inch, height=3.5*inch)
                            story.append(Spacer(1, 6))
                            story.append(img)
                            story.append(Spacer(1, 6))
                            image_index += 1
                        except Exception as e:
                            logger.warning(f"Could not add image to PDF: {e}")
                            image_index += 1
                
                # Add extra spacing for empty paragraphs (which might indicate page structure)
                elif not text and paragraph_count > 0:
                    story.append(Spacer(1, 12))
            
            # Process tables with proper spacing to maintain page flow
            for table in docx_doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else " ")
                    if any(row_data):
                        table_data.append(row_data)
                
                if table_data:
                    # Create table with proper sizing
                    pdf_table = Table(table_data, repeatRows=1)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
                    ]))
                    
                    story.append(Spacer(1, 10))
                    story.append(pdf_table)
                    story.append(Spacer(1, 10))
            
            # Add any remaining images
            while image_index < len(extracted_images):
                try:
                    img_path, _ = extracted_images[image_index]
                    img = RLImage(img_path, width=5*inch, height=3.5*inch)
                    story.append(Spacer(1, 12))
                    story.append(img)
                    story.append(Spacer(1, 12))
                    image_index += 1
                except Exception as e:
                    logger.warning(f"Could not add remaining image to PDF: {e}")
                    image_index += 1
            
            # Handle empty document case
            if not story:
                no_content_para = Paragraph(
                    f"The document '{filename}' appears to be empty or could not be processed. "
                    "Please ensure the document contains readable content.", 
                    normal_style
                )
                story = [no_content_para]
            
            # Build PDF with automatic page breaks based on content flow
            doc.build(story)
            
            # Clean up temp files
            if temp_docx_path and os.path.exists(temp_docx_path):
                os.unlink(temp_docx_path)
            
            for img_path in temp_images:
                try:
                    if os.path.exists(img_path):
                        os.unlink(img_path)
                except:
                    pass
            
            pdf_buffer.seek(0)
            pdf_content = pdf_buffer.getvalue()
            
            # Log page count for debugging
            try:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(pdf_content))
                logger.info(f"Converted '{filename}' to PDF with {len(reader.pages)} pages")
            except:
                pass
                
            return pdf_content
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            # Clean up temp files if they exist
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    os.unlink(temp_docx_path)
                except:
                    pass
                    
            for img_path in temp_images:
                try:
                    if os.path.exists(img_path):
                        os.unlink(img_path)
                except:
                    pass
                    
            raise Exception(f"Failed to convert {filename}: {str(e)}")

    @staticmethod
    def merge_pdfs(pdf_contents: List[bytes]) -> bytes:
        """Merge multiple PDF files into a single PDF"""
        try:
            writer = PdfWriter()
            
            for pdf_content in pdf_contents:
                pdf_buffer = io.BytesIO(pdf_content)
                reader = PdfReader(pdf_buffer)
                
                for page in reader.pages:
                    writer.add_page(page)
            
            output_buffer = io.BytesIO()
            writer.write(output_buffer)
            output_buffer.seek(0)
            
            return output_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error merging PDFs: {str(e)}")
            raise Exception(f"Failed to merge PDFs: {str(e)}")

    @staticmethod
    def create_zip_with_pdfs(pdf_files: List[tuple]) -> bytes:
        """Create a ZIP file containing multiple PDFs"""
        try:
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for pdf_content, original_filename in pdf_files:
                    # Change extension from .docx to .pdf
                    pdf_filename = original_filename.rsplit('.', 1)[0] + '.pdf'
                    zip_file.writestr(pdf_filename, pdf_content)
            
            zip_buffer.seek(0)
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error creating ZIP: {str(e)}")
            raise Exception(f"Failed to create ZIP: {str(e)}")

    @staticmethod
    def compress_pdf(pdf_content: bytes, compression_level: int = 90) -> bytes:
        """Compress a PDF file by reducing quality"""
        try:
            input_buffer = io.BytesIO(pdf_content)
            reader = PdfReader(input_buffer)
            writer = PdfWriter()
            
            # Add all pages to writer
            for page in reader.pages:
                writer.add_page(page)
            
            # Compress the PDF
            writer.compress_identical_objects()
            
            output_buffer = io.BytesIO()
            writer.write(output_buffer)
            output_buffer.seek(0)
            
            compressed_content = output_buffer.getvalue()
            
            # Calculate compression ratio
            original_size = len(pdf_content)
            compressed_size = len(compressed_content)
            
            logger.info(f"PDF compression: {original_size} -> {compressed_size} bytes")
            
            return compressed_content
            
        except Exception as e:
            logger.error(f"Error compressing PDF: {str(e)}")
            raise Exception(f"Failed to compress PDF: {str(e)}")