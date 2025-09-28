#!/usr/bin/env python3
"""
Backend API Testing for PDF Tools Application
Tests Word to PDF conversion, PDF compression, and error handling
"""

import requests
import io
import os
import tempfile
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import zipfile

# Get backend URL from environment
BACKEND_URL = "https://doctools-6.preview.emergentagent.com/api"

def create_test_docx():
    """Create a test DOCX file for testing"""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph for PDF conversion.')
    doc.add_paragraph('This document contains multiple paragraphs to test the conversion process.')
    doc.add_heading('Section 1', level=1)
    doc.add_paragraph('Content under section 1.')
    doc.add_heading('Section 2', level=1)
    doc.add_paragraph('Content under section 2 with some more text to make it substantial.')
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()

def create_test_pdf():
    """Create a test PDF file for compression testing"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add content to make the PDF substantial
    story.append(Paragraph("Test PDF Document", styles['Title']))
    for i in range(10):
        story.append(Paragraph(f"This is paragraph {i+1} with some content to make the PDF file larger for compression testing. " * 5, styles['Normal']))
    
    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()

def test_root_endpoint():
    """Test the root API endpoint"""
    print("🔍 Testing root endpoint...")
    try:
        response = requests.get(f"{BACKEND_URL}/")
        print(f"   Status Code: {response.status_code}")
        print(f"   Response: {response.json()}")
        
        if response.status_code == 200:
            print("   ✅ Root endpoint working")
            return True
        else:
            print("   ❌ Root endpoint failed")
            return False
    except Exception as e:
        print(f"   ❌ Root endpoint error: {str(e)}")
        return False

def test_word_to_pdf_single():
    """Test Word to PDF conversion with single_pdf=True"""
    print("\n🔍 Testing Word to PDF conversion (single PDF)...")
    try:
        docx_content = create_test_docx()
        
        files = {'files': ('test_document.docx', docx_content, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        print(f"   Status Code: {response.status_code}")
        print(f"   Content-Type: {response.headers.get('content-type')}")
        print(f"   Content-Disposition: {response.headers.get('content-disposition')}")
        print(f"   Response Size: {len(response.content)} bytes")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/pdf':
                # Verify it's a valid PDF by checking PDF header
                if response.content.startswith(b'%PDF'):
                    print("   ✅ Single PDF conversion successful - Valid PDF generated")
                    return True
                else:
                    print("   ❌ Response is not a valid PDF")
                    return False
            else:
                print("   ❌ Wrong content type returned")
                return False
        else:
            print(f"   ❌ Conversion failed: {response.text}")
            return False
    except Exception as e:
        print(f"   ❌ Word to PDF conversion error: {str(e)}")
        return False

def test_word_to_pdf_zip():
    """Test Word to PDF conversion with single_pdf=False (ZIP output)"""
    print("\n🔍 Testing Word to PDF conversion (ZIP output)...")
    try:
        docx_content1 = create_test_docx()
        docx_content2 = create_test_docx()
        
        files = [
            ('files', ('document1.docx', docx_content1, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')),
            ('files', ('document2.docx', docx_content2, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'))
        ]
        data = {'single_pdf': 'false'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        print(f"   Status Code: {response.status_code}")
        print(f"   Content-Type: {response.headers.get('content-type')}")
        print(f"   Content-Disposition: {response.headers.get('content-disposition')}")
        print(f"   Response Size: {len(response.content)} bytes")
        
        if response.status_code == 200:
            if response.headers.get('content-type') == 'application/zip':
                # Verify it's a valid ZIP by trying to open it
                try:
                    zip_buffer = io.BytesIO(response.content)
                    with zipfile.ZipFile(zip_buffer, 'r') as zip_file:
                        file_list = zip_file.namelist()
                        print(f"   ZIP contains files: {file_list}")
                        if len(file_list) == 2 and all(f.endswith('.pdf') for f in file_list):
                            print("   ✅ ZIP conversion successful - Contains expected PDF files")
                            return True
                        else:
                            print("   ❌ ZIP doesn't contain expected PDF files")
                            return False
                except Exception as zip_error:
                    print(f"   ❌ Invalid ZIP file: {str(zip_error)}")
                    return False
            else:
                print("   ❌ Wrong content type returned")
                return False
        else:
            print(f"   ❌ Conversion failed: {response.text}")
            return False
    except Exception as e:
        print(f"   ❌ Word to PDF ZIP conversion error: {str(e)}")
        return False

def test_pdf_compression():
    """Test PDF compression functionality"""
    print("\n🔍 Testing PDF compression...")
    try:
        pdf_content = create_test_pdf()
        original_size = len(pdf_content)
        print(f"   Original PDF size: {original_size} bytes")
        
        files = {'file': ('test_document.pdf', pdf_content, 'application/pdf')}
        data = {'compression_level': '90'}
        
        response = requests.post(f"{BACKEND_URL}/compress/pdf", files=files, data=data)
        print(f"   Status Code: {response.status_code}")
        print(f"   Content-Type: {response.headers.get('content-type')}")
        print(f"   Content-Disposition: {response.headers.get('content-disposition')}")
        
        if response.status_code == 200:
            compressed_size = len(response.content)
            print(f"   Compressed PDF size: {compressed_size} bytes")
            
            if response.headers.get('content-type') == 'application/pdf':
                # Verify it's a valid PDF
                if response.content.startswith(b'%PDF'):
                    compression_ratio = (1 - compressed_size / original_size) * 100
                    print(f"   Compression ratio: {compression_ratio:.1f}%")
                    print("   ✅ PDF compression successful - Valid compressed PDF generated")
                    return True
                else:
                    print("   ❌ Response is not a valid PDF")
                    return False
            else:
                print("   ❌ Wrong content type returned")
                return False
        else:
            print(f"   ❌ Compression failed: {response.text}")
            return False
    except Exception as e:
        print(f"   ❌ PDF compression error: {str(e)}")
        return False

def test_error_handling():
    """Test error handling for invalid inputs"""
    print("\n🔍 Testing error handling...")
    
    # Test 1: Invalid file type for Word conversion
    print("   Testing invalid file type for Word conversion...")
    try:
        invalid_content = b"This is not a DOCX file"
        files = {'files': ('test.txt', invalid_content, 'text/plain')}
        data = {'single_pdf': 'true'}
        
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files=files, data=data)
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code == 400:
            print("   ✅ Correctly rejected invalid file type for Word conversion")
            error_1_pass = True
        else:
            print(f"   ❌ Should have returned 400, got {response.status_code}")
            error_1_pass = False
    except Exception as e:
        print(f"   ❌ Error testing invalid Word file: {str(e)}")
        error_1_pass = False
    
    # Test 2: Invalid file type for PDF compression
    print("   Testing invalid file type for PDF compression...")
    try:
        invalid_content = b"This is not a PDF file"
        files = {'file': ('test.txt', invalid_content, 'text/plain')}
        data = {'compression_level': '90'}
        
        response = requests.post(f"{BACKEND_URL}/compress/pdf", files=files, data=data)
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code == 400:
            print("   ✅ Correctly rejected invalid file type for PDF compression")
            error_2_pass = True
        else:
            print(f"   ❌ Should have returned 400, got {response.status_code}")
            error_2_pass = False
    except Exception as e:
        print(f"   ❌ Error testing invalid PDF file: {str(e)}")
        error_2_pass = False
    
    # Test 3: Empty request for Word conversion
    print("   Testing empty request for Word conversion...")
    try:
        response = requests.post(f"{BACKEND_URL}/convert/word-to-pdf", files={}, data={'single_pdf': 'true'})
        print(f"   Status Code: {response.status_code}")
        
        if response.status_code == 422:  # FastAPI validation error
            print("   ✅ Correctly rejected empty request for Word conversion")
            error_3_pass = True
        else:
            print(f"   ❌ Should have returned 422, got {response.status_code}")
            error_3_pass = False
    except Exception as e:
        print(f"   ❌ Error testing empty Word request: {str(e)}")
        error_3_pass = False
    
    return error_1_pass and error_2_pass and error_3_pass

def test_cors_headers():
    """Test CORS headers are properly set"""
    print("\n🔍 Testing CORS headers...")
    try:
        response = requests.options(f"{BACKEND_URL}/")
        print(f"   Status Code: {response.status_code}")
        
        cors_headers = {
            'access-control-allow-origin': response.headers.get('access-control-allow-origin'),
            'access-control-allow-methods': response.headers.get('access-control-allow-methods'),
            'access-control-allow-headers': response.headers.get('access-control-allow-headers')
        }
        
        print(f"   CORS Headers: {cors_headers}")
        
        if cors_headers['access-control-allow-origin'] == '*':
            print("   ✅ CORS headers properly configured")
            return True
        else:
            print("   ❌ CORS headers not properly configured")
            return False
    except Exception as e:
        print(f"   ❌ CORS test error: {str(e)}")
        return False

def main():
    """Run all tests"""
    print("🚀 Starting PDF Tools Backend API Tests")
    print(f"Backend URL: {BACKEND_URL}")
    print("=" * 60)
    
    test_results = {}
    
    # Run all tests
    test_results['root_endpoint'] = test_root_endpoint()
    test_results['word_to_pdf_single'] = test_word_to_pdf_single()
    test_results['word_to_pdf_zip'] = test_word_to_pdf_zip()
    test_results['pdf_compression'] = test_pdf_compression()
    test_results['error_handling'] = test_error_handling()
    test_results['cors_headers'] = test_cors_headers()
    
    # Summary
    print("\n" + "=" * 60)
    print("📊 TEST SUMMARY")
    print("=" * 60)
    
    passed = 0
    total = len(test_results)
    
    for test_name, result in test_results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{test_name.replace('_', ' ').title()}: {status}")
        if result:
            passed += 1
    
    print(f"\nOverall: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All tests passed! PDF Tools backend is working correctly.")
        return True
    else:
        print("⚠️  Some tests failed. Please check the issues above.")
        return False

if __name__ == "__main__":
    main()